from docx import Document
from bs4 import BeautifulSoup
import requests
import re 
class Wiki:
    url=str()
    stp=str()
    def __init__(self,kw,docu): #kw stands for keyword
        self.url="https://en.wikipedia.org/wiki/"+kw.replace(" ","_") 
        self.doc=BeautifulSoup(requests.get(self.url).text,"html.parser")
        self.docu=docu
        docu.add_heading(kw,0)
        docu.add_heading(self.url,1)
    
    def trip(self):
        self.q=self.doc.find_all(["p","h3","h2","ul"])
        for i in self.q:
            if(i.name=="h2" and list(i.strings)[0]=='See also'):
                break
            #if(i.string=="Contents"):
                #continue
            if(i.name=="p"):
                k=True
                for j in i.strings:
                    if(j[0]=='[' and j[-1]==']'):
                        continue
                    elif(j=="["):
                        k=False
                        continue
                    elif(j==']'):
                        k=True
                        continue
                    if(k==True):
                        self.stp=self.stp + str(j)
                docu.add_paragraph(self.stp)
                self.stp=''
            elif(i.name=="h3"):
                k=True
                for j in i.strings:
                    if(j[0]=='[' and j[-1]==']'):
                        continue
                    elif(j=="["):
                        k=False
                        continue
                    elif(j==']'):
                        k=True
                        continue
                    if(k==True):
                        self.stp=self.stp + str(j)
                docu.add_heading(self.stp,3)
                self.stp=''
            elif(i.name=="h2"):
                k=True
                for j in i.strings:
                    if(j[0]=='[' and j[-1]==']'):
                        continue
                    elif(j=="["):
                        k=False
                        continue
                    elif(j==']'):
                        k=True
                        continue
                    if(k==True):
                        self.stp=self.stp + str(j)
                docu.add_heading(self.stp,2)
                self.stp=''
            elif(i.name=='ul'):
                k=True
                for j in i.strings:
                    if(j[0]=='[' and j[-1]==']'):
                        continue
                    elif(j=="["):
                        k=False
                        continue
                    elif(j==']'):
                        k=True
                        continue
                    if(k==True):
                        self.stp=self.stp + str(j)
                docu.add_paragraph(self.stp)
                self.stp=''



docu=Document()
ij=input("PLease enter the key word : ")
j=Wiki(ij,docu)
j.trip()
print("Code sucess")
docu.save("rp.docx")

